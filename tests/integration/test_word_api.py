"""集成测试：Word 导出/导入 API 端点。"""
from __future__ import annotations

import json
from io import BytesIO

import pytest
from docx import Document

pytestmark = pytest.mark.asyncio


# =========================
# 辅助函数
# =========================


async def _create_student(api_client, name="张三"):
    """创建学生并返回 ID。"""
    r = await api_client.post(
        "/api/students",
        json={"name": name, "grade": "三年级", "age": 10},
    )
    assert r.status_code == 201, r.text
    return r.json()["id"]


async def _create_record(api_client, student_id):
    """创建课程记录并返回 ID。"""
    r = await api_client.post(
        "/api/reports",
        json={
            "student_id": student_id,
            "course_date": "2025-06-15",
            "course_topic": "Python 入门",
            "knowledge_points": ["变量", "循环"],
            "ability_improvement": "逻辑思维提升",
            "evaluation": "学生表现优秀",
            "homework": {"goal": "完成计算器", "hints": [], "criteria": []},
            "vocabulary": {"word": "variable", "meaning": "变量"},
            "content_items": [{"kp": "变量", "text": "学习变量定义"}],
        },
    )
    assert r.status_code == 201, r.text
    return r.json()["id"]


# =========================
# 测试：Word 导出
# =========================


class TestExportWord:
    async def test_export_word_success(self, api_client):
        """导出 Word 应返回 docx_path。"""
        sid = await _create_student(api_client)
        rid = await _create_record(api_client, sid)

        r = await api_client.post(f"/api/reports/{rid}/export-word", json={"template_id": "classic"})
        assert r.status_code == 200, r.text
        data = r.json()
        assert "docx_path" in data
        assert data["docx_path"].endswith(".docx")
        assert "filename" in data

    async def test_export_word_stream(self, api_client):
        """流式导出 Word 应返回 docx 内容。"""
        sid = await _create_student(api_client)
        rid = await _create_record(api_client, sid)

        r = await api_client.post(
            f"/api/reports/{rid}/export-word-stream", json={"template_id": "classic"}
        )
        assert r.status_code == 200, r.text
        assert (
            r.headers["content-type"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        content = r.content
        assert len(content) > 0
        # 验证是有效 docx
        doc = Document(BytesIO(content))
        assert doc is not None

    async def test_export_word_nonexistent_record(self, api_client):
        """不存在的记录应返回 404。"""
        r = await api_client.post("/api/reports/99999/export-word", json={"template_id": "classic"})
        assert r.status_code == 404

    async def test_export_word_contains_student_name(self, api_client):
        """导出的 Word 应包含学生姓名。"""
        sid = await _create_student(api_client, "李四")
        rid = await _create_record(api_client, sid)

        r = await api_client.post(f"/api/reports/{rid}/export-word-stream", json={"template_id": "classic"})
        assert r.status_code == 200
        doc = Document(BytesIO(r.content))
        texts = [p.text for p in doc.paragraphs]
        assert any("李四" in t for t in texts)

    async def test_export_word_contains_date(self, api_client):
        """导出的 Word 应包含课程日期。"""
        sid = await _create_student(api_client)
        rid = await _create_record(api_client, sid)

        r = await api_client.post(f"/api/reports/{rid}/export-word-stream", json={"template_id": "classic"})
        assert r.status_code == 200
        doc = Document(BytesIO(r.content))
        texts = [p.text for p in doc.paragraphs]
        assert any("2025-06-15" in t for t in texts)

    async def test_export_word_with_different_template(self, api_client):
        """使用不同模板应仍可生成 docx。"""
        sid = await _create_student(api_client)
        rid = await _create_record(api_client, sid)

        r = await api_client.post(f"/api/reports/{rid}/export-word", json={"template_id": "cartoon"})
        assert r.status_code == 200, r.text
        data = r.json()
        assert "docx_path" in data


# =========================
# 测试：Word 导入
# =========================


class TestImportWord:
    async def test_import_word_success(self, api_client):
        """导入有效 docx 应解析出字段。"""
        # 先导出一个，再导入
        sid = await _create_student(api_client)
        rid = await _create_record(api_client, sid)

        export_r = await api_client.post(
            f"/api/reports/{rid}/export-word-stream", json={"template_id": "classic"}
        )
        assert export_r.status_code == 200
        docx_bytes = export_r.content

        # 导入
        r = await api_client.post(
            "/api/import/word",
            files={
                "file": (
                    "test.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert r.status_code == 200, r.text
        data = r.json()
        assert "fields" in data
        assert "unrecognized_sections" in data
        assert "confidence" in data

    async def test_import_word_roundtrip(self, api_client):
        """导出再导入应还原大部分字段。"""
        sid = await _create_student(api_client, "王五")
        rid = await _create_record(api_client, sid)

        # 导出
        export_r = await api_client.post(
            f"/api/reports/{rid}/export-word-stream", json={"template_id": "classic"}
        )
        assert export_r.status_code == 200

        # 导入
        r = await api_client.post(
            "/api/import/word",
            files={
                "file": (
                    "roundtrip.docx",
                    export_r.content,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert r.status_code == 200, r.text
        fields = r.json()["fields"]

        # 验证关键字段
        assert fields.get("course_topic") == "Python 入门", str(fields)
        assert fields.get("course_date") == "2025-06-15"

    async def test_import_invalid_format(self, api_client):
        """非 docx 文件应拒绝。"""
        r = await api_client.post(
            "/api/import/word",
            files={"file": ("test.txt", b"hello world", "text/plain")},
        )
        assert r.status_code == 400

    async def test_import_empty_file(self, api_client):
        """空文件应返回错误。"""
        r = await api_client.post(
            "/api/import/word",
            files={
                "file": (
                    "empty.docx",
                    b"",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert r.status_code == 400

    async def test_import_no_file(self, api_client):
        """不传文件应返回 422。"""
        r = await api_client.post("/api/import/word")
        assert r.status_code == 422
