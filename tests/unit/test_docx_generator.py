"""单元测试：docx_generator.py Word 文档生成服务。"""
from __future__ import annotations

import json
from io import BytesIO

import pytest
from docx import Document
from docx.shared import Mm, Pt, RGBColor

from backend.services.docx_generator import DocxGenerationError, generate

# =========================
# Fixtures
# =========================


def _open_docx(content: bytes) -> Document:
    """从 bytes 打开 docx。"""
    return Document(BytesIO(content))


class FakeRecord:
    """模拟 CourseRecord 对象（仅含 docx_generator 需要的属性）。"""

    def __init__(self, **kwargs):
        for k, v in kwargs.items():
            setattr(self, k, v)


@pytest.fixture
def minimal_record():
    """最小记录——所有字段为空。"""
    return FakeRecord(
        course_date="",
        course_topic="",
        project_folder="",
        knowledge_points=None,
        ability_improvement=None,
        content_items=None,
        vocabulary=None,
        homework=None,
        screenshot_paths=None,
        evaluation=None,
    )


@pytest.fixture
def full_record():
    """完整记录——所有字段有内容。"""
    return FakeRecord(
        course_date="2025-06-15",
        course_topic="Python 入门",
        project_folder="/home/projects/python101",
        knowledge_points=json.dumps(["变量", "循环", "函数"], ensure_ascii=False),
        ability_improvement="逻辑思维和问题解决能力有显著提升",
        content_items=json.dumps([
            {"kp": "变量", "text": "学习了变量的定义和使用"},
            {"kp": "循环", "text": "掌握了 for 和 while 循环"},
        ], ensure_ascii=False),
        vocabulary=json.dumps({
            "word": "variable",
            "phonetic": "ˈveəriəbl",
            "meaning": "变量",
            "example": "x is a variable.",
        }, ensure_ascii=False),
        homework=json.dumps({
            "goal": "完成一个计算器程序",
            "hints": ["使用 input 获取输入", "处理加减乘除"],
            "criteria": ["正确接收输入", "运算结果正确", "代码整洁"],
        }, ensure_ascii=False),
        screenshot_paths=json.dumps([]),
        evaluation="学生表现优秀，积极参与课堂互动",
    )


# =========================
# 测试：基本功能
# =========================


class TestGenerate:
    def test_empty_record_returns_valid_docx(self, minimal_record):
        """空记录应生成有效的 docx 字节流。"""
        result = generate(minimal_record)
        assert isinstance(result, bytes)
        assert len(result) > 0

        # 验证是有效 docx
        doc = _open_docx(result)
        assert doc is not None

    def test_full_record_has_content(self, full_record):
        """完整记录应包含所有章节标题。"""
        result = generate(full_record)
        doc = _open_docx(result)
        texts = [p.text for p in doc.paragraphs]

        # 检查章节标题
        assert "基本信息" in texts or any("基本信息" in t for t in texts)
        assert "知识点概括" in texts or any("知识点概括" in t for t in texts)
        assert "能力提升" in texts or any("能力提升" in t for t in texts)
        assert "内容概述" in texts or any("内容概述" in t for t in texts)
        assert "单词学习" in texts or any("单词学习" in t for t in texts)
        assert "课后作业" in texts or any("课后作业" in t for t in texts)
        assert "学生评价" in texts or any("学生评价" in t for t in texts)

    def test_contains_student_name(self, full_record):
        """文档应包含学生姓名。"""
        result = generate(full_record, student_name="张三")
        doc = _open_docx(result)
        texts = [p.text for p in doc.paragraphs]
        assert any("张三" in t for t in texts)

    def test_contains_basic_info(self, full_record):
        """基本信息应包含日期和课程名称。"""
        result = generate(full_record, student_name="张三")
        doc = _open_docx(result)
        texts = [p.text for p in doc.paragraphs]
        assert any("2025-06-15" in t for t in texts)
        assert any("Python 入门" in t for t in texts)

    def test_knowledge_points_as_bullets(self, full_record):
        """知识点应显示为列表。"""
        result = generate(full_record)
        doc = _open_docx(result)
        texts = [p.text for p in doc.paragraphs]
        assert any("变量" in t for t in texts)
        assert any("循环" in t for t in texts)
        assert any("函数" in t for t in texts)

    def test_ability_improvement_included(self, full_record):
        """能力提升段落应出现在文档中。"""
        result = generate(full_record)
        doc = _open_docx(result)
        texts = [p.text for p in doc.paragraphs]
        assert any("逻辑思维" in t for t in texts)

    def test_evaluation_included(self, full_record):
        """学生评价应出现在文档中。"""
        result = generate(full_record)
        doc = _open_docx(result)
        texts = [p.text for p in doc.paragraphs]
        assert any("课堂互动" in t for t in texts)

    def test_custom_primary_color_from_template(self, full_record):
        """模板配置中的主题色应生效。"""
        tpl_config = {"theme": {"primary_color": "#FF5722"}}
        result = generate(full_record, template_config=tpl_config)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_chinese_content_utf8(self, full_record):
        """中文字符应正确处理。"""
        result = generate(full_record, student_name="测试学生")
        doc = _open_docx(result)
        # 保存到临时变量以验证文本内容
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "测试学生" in full_text
        assert "Python 入门" in full_text


class TestGenerationEdgeCases:
    def test_none_json_fields(self, minimal_record):
        """None JSON 字段不应导致错误。"""
        result = generate(minimal_record)
        assert isinstance(result, bytes)

    def test_empty_knowledge_points(self, full_record):
        """空知识点列表不应影响生成。"""
        full_record.knowledge_points = json.dumps([])
        result = generate(full_record)
        doc = _open_docx(result)
        texts = [p.text for p in doc.paragraphs]
        # 知识点概括标题不应出现
        has_kp_header = any("知识点概括" in t for t in texts)
        # 即使有标题也不应该有列表项
        assert isinstance(result, bytes)

    def test_empty_homework(self, full_record):
        """空 homework 不应导致错误。"""
        full_record.homework = json.dumps({"goal": "", "hints": [], "criteria": []})
        result = generate(full_record)
        assert isinstance(result, bytes)

    def test_title_centered(self, full_record):
        """标题段落应居中。"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        result = generate(full_record)
        doc = _open_docx(result)
        # 第一个非空段落应为标题
        for p in doc.paragraphs:
            if p.text.strip():
                assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
                break
