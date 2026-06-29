"""单元测试：docx_importer.py Word 文档导入解析服务。"""
from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from docx.shared import Pt

from backend.services.docx_importer import DocxImportError, import_docx


# =========================
# Fixtures
# =========================


def _make_docx(headings_and_paras: list[tuple[str, list[str]]]) -> bytes:
    """生成带指定 Heading 1 分段和内容的 docx bytes。"""
    doc = Document()
    for heading_text, paragraphs in headings_and_paras:
        doc.add_heading(heading_text, level=1)
        for p_text in paragraphs:
            doc.add_paragraph(p_text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def full_docx():
    """完整的课程报告 docx。"""
    return _make_docx([
        ("基本信息", [
            "上课时间：2025-06-15",
            "学生姓名：张三",
            "课程名称：Python 入门",
        ]),
        ("知识点概括", [
            "变量",
            "循环",
            "函数",
        ]),
        ("能力提升", [
            "逻辑思维和问题解决能力有显著提升",
        ]),
        ("内容概述", [
            "变量：学习了变量的定义和使用",
            "循环：掌握了 for 和 while 循环",
        ]),
        ("单词学习", [
            "单词：variable",
            "音标：/ˈveəriəbl/",
            "释义：变量",
            "例句：x is a variable.",
        ]),
        ("课后作业", [
            "目标：完成一个计算器程序",
            "提示",
            "使用 input 获取输入",
            "处理加减乘除",
            "评分点",
            "正确接收输入",
            "运算结果正确",
        ]),
        ("学生评价", [
            "学生表现优秀，积极参与课堂互动",
        ]),
    ])


@pytest.fixture
def minimal_docx():
    """只含基本信息的 docx。"""
    return _make_docx([
        ("基本信息", [
            "上课时间：2025-06-15",
            "课程名称：测试课程",
        ]),
    ])


# =========================
# 测试：基本导入
# =========================


class TestImportDocx:
    def test_full_docx_parsed(self, full_docx):
        """完整 docx 应解析出所有字段。"""
        result = import_docx(full_docx)
        fields = result["fields"]
        assert "course_date" in fields
        assert "course_topic" in fields
        assert "knowledge_points" in fields
        assert "ability_improvement" in fields
        assert "content_items" in fields
        assert "vocabulary" in fields
        assert "homework" in fields
        assert "evaluation" in fields

    def test_basic_info_parsed(self, full_docx):
        """基本信息应正确解析。"""
        result = import_docx(full_docx)
        fields = result["fields"]
        assert fields["course_date"] == "2025-06-15"
        assert fields["course_topic"] == "Python 入门"

    def test_knowledge_points(self, full_docx):
        """知识点应解析为列表。"""
        result = import_docx(full_docx)
        kp = result["fields"].get("knowledge_points", [])
        assert isinstance(kp, list)
        assert "变量" in kp
        assert "循环" in kp
        assert "函数" in kp

    def test_content_items(self, full_docx):
        """内容概述应解析为 [{kp, text}]。"""
        result = import_docx(full_docx)
        items = result["fields"].get("content_items", [])
        assert isinstance(items, list)
        assert len(items) >= 2
        assert items[0]["kp"] == "变量"
        assert "变量" in items[0]["text"]

    def test_vocabulary(self, full_docx):
        """单词学习应正确解析。"""
        result = import_docx(full_docx)
        vocab = result["fields"].get("vocabulary", {})
        assert vocab.get("word") == "variable"
        assert "ˈveəriəbl" in vocab.get("phonetic", "")
        assert vocab.get("meaning") == "变量"

    def test_homework_goal(self, full_docx):
        """课后作业应解析出目标和提示。"""
        result = import_docx(full_docx)
        hw = result["fields"].get("homework", {})
        assert "计算器程序" in hw.get("goal", "")
        assert len(hw.get("hints", [])) >= 1
        assert len(hw.get("criteria", [])) >= 1

    def test_evaluation(self, full_docx):
        """学生评价应正确解析。"""
        result = import_docx(full_docx)
        ev = result["fields"].get("evaluation", "")
        assert "课堂互动" in ev

    def test_minimal_docx(self, minimal_docx):
        """最小 docx 应解析基本信息。"""
        result = import_docx(minimal_docx)
        fields = result["fields"]
        assert fields.get("course_date") == "2025-06-15"
        assert fields.get("course_topic") == "测试课程"

    def test_confidence_score(self, full_docx):
        """完整文档应有较高置信度。"""
        result = import_docx(full_docx)
        assert result["confidence"] > 0.5

    def test_unrecognized_sections_empty(self, full_docx):
        """完整文档应无未识别小节。"""
        result = import_docx(full_docx)
        assert len(result["unrecognized_sections"]) == 0

    def test_result_structure(self, full_docx):
        """返回结构应包含 fields / unrecognized_sections / confidence。"""
        result = import_docx(full_docx)
        assert "fields" in result
        assert "unrecognized_sections" in result
        assert "confidence" in result


class TestImportEdgeCases:
    def test_invalid_bytes(self):
        """无效字节应抛出 DocxImportError。"""
        with pytest.raises(DocxImportError):
            import_docx(b"not a docx file")

    def test_empty_bytes(self):
        """空字节应抛出 DocxImportError。"""
        with pytest.raises(DocxImportError):
            import_docx(b"")

    def test_no_heading_docx(self):
        """无 Heading 1 的文档应只含基本信息。"""
        doc = Document()
        doc.add_paragraph("一些文本")
        doc.add_paragraph("更多文本")
        buf = BytesIO()
        doc.save(buf)
        result = import_docx(buf.getvalue())
        # 无匹配小节
        assert result["unrecognized_sections"] == []
        assert result["confidence"] == 0

    def test_unknown_section(self):
        """未知标题应归入 unrecognized_sections。"""
        docx = _make_docx([
            ("自定义未知章节", ["这是不认识的文本"]),
        ])
        result = import_docx(docx)
        assert len(result["unrecognized_sections"]) >= 1
        assert result["confidence"] == 0

    def test_ability_improvement_parsed(self, full_docx):
        """能力提升应解析。"""
        result = import_docx(full_docx)
        ability = result["fields"].get("ability_improvement", "")
        assert "逻辑思维" in ability

    def test_docx_with_tables(self):
        """含表格的 docx 不应崩溃（仅忽略表格）。"""
        doc = Document()
        doc.add_heading("基本信息", level=1)
        doc.add_paragraph("上课时间：2025-06-15")
        doc.add_heading("课后作业", level=1)
        table = doc.add_table(2, 2)
        table.cell(0, 0).text = "目标"
        table.cell(0, 1).text = "完成作业"
        buf = BytesIO()
        doc.save(buf)
        result = import_docx(buf.getvalue())
        assert "course_date" in result["fields"]
