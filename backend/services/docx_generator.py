"""Word 文档生成服务：CourseRecord → .docx。

使用 python-docx 原生生成，结构尽量与 PDF 模板（classic）一致：
- 4 页结构 + 分页符
- 知识点卡片（表格模拟）
- 截图嵌入（首页）
- 内容概述分页（第 2 页 2 条，后续在第 3 页）
- 代码片段（第 2 页）
- 词汇卡（第 3 页）
- 课后作业 + 学生评价（第 4 页）
"""
from __future__ import annotations

import json
import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Cm

from backend.config import PROJECT_ROOT, get_settings
from backend.utils.logger import get_logger

log = get_logger(__name__)


class DocxGenerationError(Exception):
    def __init__(self, message: str, original: Exception | None = None):
        self.original = original
        super().__init__(message)


def _load_json(value: str | None, default: Any = None) -> Any:
    if value is None:
        return default
    try:
        return json.loads(value)
    except (json.JSONDecodeError, TypeError):
        return default


def _url_to_fs_path(url_path: str) -> str | None:
    settings = get_settings()
    if url_path.startswith("/api/assets/screenshots/"):
        filename = url_path[len("/api/assets/screenshots/") :]
        full = Path(settings.report.screenshot_dir) / filename
        if full.exists():
            return str(full.resolve())
    elif url_path.startswith("/api/assets/"):
        filename = url_path[len("/api/assets/") :]
        full = Path(settings.report.asset_dir) / filename
        if full.exists():
            return str(full.resolve())
    return None


# ── 字体/样式工具 ──────────────────────────────────────────

def _set_run_font(run, font_name: str = "宋体", size: int = 11,
                  bold: bool = False, color: str | None = None,
                  font_name_ascii: str | None = None) -> None:
    """设置 run 的字体属性，中英文分别指定。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name_ascii or font_name
    # 中文字体回退
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    if font_name_ascii:
        rFonts.set(qn("w:ascii"), font_name_ascii)
        rFonts.set(qn("w:hAnsi"), font_name_ascii)
    if color:
        try:
            run.font.color.rgb = RGBColor(
                int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
            )
        except (ValueError, IndexError):
            pass


def _add_heading(doc, text: str, font_name: str, font_size: int,
                 color: str) -> None:
    """添加带渐变色带的标题行（类似模板的 section-title）。"""
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(4)

    # 左侧色带
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)

    r = p.add_run(text)
    _set_run_font(r, font_name, size=font_size, bold=True, color=color)
    return p


def _add_body(doc, text: str, font_name: str, font_size: int) -> None:
    p = doc.add_paragraph(text)
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    for run in p.runs:
        _set_run_font(run, font_name, size=font_size)


def _add_labeled(doc, label: str, value: str, font_name: str,
                 font_size: int) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    r1 = p.add_run(f"{label}：")
    _set_run_font(r1, font_name, size=font_size, bold=True)
    r2 = p.add_run(value or "")
    _set_run_font(r2, font_name, size=font_size)


def _add_page_header(doc, text: str, font_name: str, font_size: int,
                     color: str) -> None:
    """添加页面头部（如"课程内容详解"）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(4)
    p.space_after = Pt(8)
    # 底部边框
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)

    r = p.add_run(text)
    _set_run_font(r, font_name, size=max(14, font_size + 2), bold=True, color=color)


def _add_page_break(doc) -> None:
    """添加分页符。"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def _add_kp_card(doc, text: str, font_name: str, font_size: int,
                 color: str) -> None:
    """知识点卡片（用带圆角色块效果的段落）。"""
    p = doc.add_paragraph()
    p.space_before = Pt(3)
    p.space_after = Pt(3)
    # 背景色阴影效果：用左侧色带 + 缩进
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)
    pPr.append(OxmlElement("w:shd"))  # 浅色背景
    shd = pPr.find(qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _lighten_hex(color, 0.85))

    # 缩进
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "284")  # ≈ 5mm
    ind.set(qn("w:hanging"), "0")
    pPr.append(ind)

    r = p.add_run(f"  {text}")
    _set_run_font(r, font_name, size=font_size, color=_lighten_hex(color, 0.3))


def _lighten_hex(hex_color: str, factor: float = 0.7) -> str:
    """将 hex 颜色变浅。factor=0 全黑, 1 全白。"""
    try:
        h = hex_color.lstrip("#")
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        r = min(255, int(r + (255 - r) * factor))
        g = min(255, int(g + (255 - g) * factor))
        b = min(255, int(b + (255 - b) * factor))
        return f"#{r:02x}{g:02x}{b:02x}"
    except (ValueError, IndexError):
        return "#f0f0f0"


def _add_code_block(doc, code: str, font_name: str, font_size: int) -> None:
    """添加代码块（灰色背景段落）。"""
    for line in code.split("\n"):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        pPr = p._element.get_or_add_pPr()
        # 灰色背景
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)
        # 等宽字体缩进
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "284")
        pPr.append(ind)

        r = p.add_run(line if line else " ")
        _set_run_font(r, "Courier New", size=font_size - 1, font_name_ascii="Courier New")


def _add_vocabulary_card(doc, vocab: dict, font_name: str, font_size: int,
                         color: str) -> None:
    """词汇卡片（带边框的表格）。"""
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(14)

    data = [
        ("单词", f"{vocab.get('word', '')}  {vocab.get('phonetic', '')}"),
        ("释义", vocab.get("meaning", "")),
        ("例句", vocab.get("example", "")),
    ]
    for i, (label, value) in enumerate(data):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        # 标签
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(label)
        _set_run_font(r0, font_name, size=font_size, bold=True, color=color)
        # 值
        p1 = cell1.paragraphs[0]
        r1 = p1.add_run(value)
        _set_run_font(r1, font_name, size=font_size)


def merge_layout_with_theme(
    template_config: dict | None,
    layout_config: dict | None,
) -> dict:
    """将布局覆盖与模板主题合并（与 report_renderer 逻辑一致）。"""
    theme = (template_config or {}).get("theme", {})
    result = {
        "primary_color": theme.get("primary_color", "#2B5FC3"),
        "secondary_color": theme.get("secondary_color", "#F0F4FA"),
        "font_title": theme.get("font_title", "Heiti SC"),
        "font_body": theme.get("font_body", "STSong"),
        "font_size_title": theme.get("font_size_title", 24),
        "font_size_body": theme.get("font_size_body", 11),
        "page_margin_top": 20,
        "page_margin_bottom": 18,
        "page_margin_left": 18,
        "page_margin_right": 18,
    }
    if layout_config:
        for key in result:
            if key in layout_config and layout_config[key] is not None:
                result[key] = layout_config[key]
    return result


def generate(
    record,
    student_name: str = "",
    template_config: dict | None = None,
    layout_config: dict | None = None,
) -> bytes:
    """生成 Word 文档字节流，结构与模板保持一致。"""
    merged = merge_layout_with_theme(template_config, layout_config)
    doc = Document()

    # ── 页面设置（A4 纵向） ──
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(merged["page_margin_top"])
    section.bottom_margin = Mm(merged["page_margin_bottom"])
    section.left_margin = Mm(merged["page_margin_left"])
    section.right_margin = Mm(merged["page_margin_right"])

    fb = merged["font_body"]
    fbs = merged["font_size_body"]
    ft = merged["font_title"]
    fts = merged["font_size_title"]
    pc = merged["primary_color"]

    # 默认样式
    style = doc.styles["Normal"]
    style.font.name = fb
    style.font.size = Pt(fbs)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), fb)

    # ── 反序列化 ──
    kp = _load_json(getattr(record, "knowledge_points", None), [])
    content_items = _load_json(getattr(record, "content_items", None), [])
    vocabulary = _load_json(getattr(record, "vocabulary", None), {})
    homework = _load_json(getattr(record, "homework", None), {})
    screenshots = _load_json(getattr(record, "screenshot_paths", None), [])

    # ── 代码片段 ──
    code_excerpt = ""
    try:
        project_meta = _load_json(getattr(record, "project_meta", None), {})
        ai_meta = _load_json(getattr(record, "ai_meta", None), {})
        key_funcs = ai_meta.get("key_functions", [])
        if key_funcs and project_meta:
            folder = project_meta.get("folder", "")
            blocks = []
            for func in key_funcs[:2]:
                fp = func.get("file_path", "")
                sl = func.get("start_line")
                el = func.get("end_line")
                if fp and sl and el and folder:
                    fpath = Path(folder) / fp
                    if fpath.exists():
                        lines = fpath.read_text("utf-8").splitlines()
                        selected = lines[sl - 1 : el]
                        blocks.append(f"# {fp}:{sl}\n" + "\n".join(selected))
            if blocks:
                code_excerpt = "\n\n".join(blocks)
    except Exception:
        pass

    # ════════════════════════════════════════════════════════
    # 第 1 页：封面 + 概览
    # ════════════════════════════════════════════════════════

    # 大标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.space_before = Pt(24)
    title_p.space_after = Pt(6)
    r = title_p.add_run(getattr(record, "course_topic", "") or "课程报告")
    _set_run_font(r, ft, size=max(22, fts), bold=True, color=pc)

    # 学生信息
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.space_after = Pt(12)
    info_text = f"{student_name} ｜ {getattr(record, 'course_date', '') or ''}"
    r = info_p.add_run(info_text)
    _set_run_font(r, fb, size=fbs, color="#666666")

    # 截图（首页展示，最多 2 张并排）
    if screenshots:
        screenshot_fs_paths = []
        for s in screenshots[:2]:
            if isinstance(s, str):
                fp = _url_to_fs_path(s)
                if fp:
                    screenshot_fs_paths.append(fp)
        if screenshot_fs_paths:
            img_table = doc.add_table(rows=1, cols=len(screenshot_fs_paths))
            img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, fp in enumerate(screenshot_fs_paths):
                cell = img_table.rows[0].cells[i]
                try:
                    run_p = cell.paragraphs[0]
                    run_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = run_p.add_run()
                    run.add_picture(fp, width=Mm(70))
                except Exception as e:
                    log.warning("截图嵌入失败 %s: %s", fp, e)
            doc.add_paragraph()  # 空行

    # 知识点
    if kp:
        _add_heading(doc, "🎯 本节课知识点", ft, max(13, fbs + 1), pc)
        for point in kp:
            _add_kp_card(doc, str(point), fb, fbs, pc)

    # 能力提升
    ability = getattr(record, "ability_improvement", "") or ""
    if ability:
        _add_heading(doc, "💪 能力提升", ft, max(13, fbs + 1), pc)
        _add_body(doc, ability, fb, fbs)

    # ════════════════════════════════════════════════════════
    # 第 2 页：内容详解 + 代码展示
    # ════════════════════════════════════════════════════════
    _add_page_break(doc)
    _add_page_header(doc, "课程内容详解", ft, fbs, pc)

    for item in content_items[:2]:
        kp_name = item.get("kp", "")
        text = item.get("text", "")
        _add_heading(doc, kp_name, ft, max(12, fbs), pc)
        _add_body(doc, text, fb, fbs)

    if code_excerpt:
        _add_heading(doc, "💻 本节课代码展示", ft, max(13, fbs + 1), pc)
        _add_code_block(doc, code_excerpt, fb, fbs)

    # ════════════════════════════════════════════════════════
    # 第 3 页：剩余内容 + 单词学习
    # ════════════════════════════════════════════════════════
    _add_page_break(doc)
    _add_page_header(doc, "课程内容（续）& 单词学习", ft, fbs, pc)

    if len(content_items) > 2:
        for item in content_items[2:]:
            kp_name = item.get("kp", "")
            text = item.get("text", "")
            _add_heading(doc, kp_name, ft, max(12, fbs), pc)
            _add_body(doc, text, fb, fbs)
    else:
        _add_body(doc, "本课知识点详解已在前页完整展示。", fb, fbs)

    if vocabulary and vocabulary.get("word"):
        _add_heading(doc, "📖 核心词汇", ft, max(13, fbs + 1), pc)
        _add_vocabulary_card(doc, vocabulary, fb, fbs, pc)

    # ════════════════════════════════════════════════════════
    # 第 4 页：作业 + 评价
    # ════════════════════════════════════════════════════════
    _add_page_break(doc)

    if homework and homework.get("goal"):
        _add_heading(doc, "📚 课后作业", ft, max(13, fbs + 1), pc)
        _add_labeled(doc, "目标", homework.get("goal", ""), fb, fbs)

        hints = homework.get("hints", [])
        if hints:
            _add_body(doc, "💡 提示：", fb, fbs)
            for h in hints:
                p = doc.add_paragraph(h, style="List Bullet")
                for run in p.runs:
                    _set_run_font(run, fb, size=fbs)

        criteria = homework.get("criteria", [])
        if criteria:
            _add_body(doc, "✅ 评分标准：", fb, fbs)
            for c in criteria:
                p = doc.add_paragraph(c, style="List Bullet")
                for run in p.runs:
                    _set_run_font(run, fb, size=fbs)

    # 学生评价
    evaluation = getattr(record, "evaluation", "") or ""
    if evaluation:
        _add_heading(doc, "⭐ 学生评价", ft, max(13, fbs + 1), pc)
        _add_body(doc, evaluation, fb, fbs)

    # ── 保存 ──
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
